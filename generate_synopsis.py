import os
import sys
import subprocess

# Ensure python-docx is installed
try:
    import docx
except ImportError:
    print("Installing python-docx...")
    subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
    import docx

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def create_synopsis():
    doc = Document()
    
    # Page margins: Standard 1 inch
    sections = doc.sections
    for section in sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Style helper
    def set_font(run, font_name="Times New Roman", size=12, bold=False, italic=False):
        run.font.name = font_name
        run.font.size = Pt(size)
        run.bold = bold
        run.italic = italic
        # Force Word to apply font name
        rPr = run._r.get_or_add_rPr()
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), font_name)
        rFonts.set(qn('w:hAnsi'), font_name)
        rPr.append(rFonts)

    def add_heading_1(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        set_font(run, size=16, bold=True)
        return p

    def add_heading_2(text):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        set_font(run, size=14, bold=True)
        return p

    def add_body(text, bold_prefix=None, italic=False):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = 1.15
        if bold_prefix:
            run_p = p.add_run(bold_prefix)
            set_font(run_p, size=12, bold=True)
        run = p.add_run(text)
        set_font(run, size=12, italic=italic)
        return p

    # --- Title Section (Properly Spaced & Formatted) ---
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_title.paragraph_format.space_before = Pt(12)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run("PROJECT SYNOPSIS\n")
    set_font(run_title, size=16, bold=True)
    
    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(24)
    run_sub = p_sub.add_run("FLOODSENSE MUMBAI & KONKAN: ML-BASED PRECIPITATION SEVERITY AND REGIONAL WATERLOGGING TELEMETRY DASHBOARD")
    set_font(run_sub, size=14, bold=True)

    # --- Student & Guide Metadata Block (Clean, bold labeled paragraphs) ---
    p_meta = doc.add_paragraph()
    p_meta.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p_meta.paragraph_format.space_after = Pt(24)
    
    metadata = [
        ("Candidate Name:   ", "Riddhi Shetye"),
        ("Roll No:          ", "260163"),
        ("Class:            ", "TYDS (Third Year Data Science)"),
        ("Project Guide:    ", "Prof. Swati Singh")
    ]
    
    for label, val in metadata:
        r_lbl = p_meta.add_run(label)
        set_font(r_lbl, size=12, bold=True)
        r_val = p_meta.add_run(val + "\n")
        set_font(r_val, size=12)

    # --- Content Sections (Justified paragraphs, 12pt Times New Roman) ---
    add_heading_1("1. Introduction & Background")
    add_body(
        "The coastal belt of Maharashtra, encompassing the metropolitan city of Mumbai and the wider Konkan geographical division, experiences heavy orographic precipitation during the South-West monsoons. While Mumbai struggles with dense urban stormwater flooding caused by concrete imperviousness and high-tide drainage blocks, the broader Konkan region faces rapid flash flooding in river basins and low-lying agricultural plains. Traditional flood-monitoring setups are highly reactive, recording waterlogging only after the damage has occurred. This project introduces FloodSense Mumbai & Konkan, an intelligent analytical framework that leverages machine learning to forecast flood severity levels proactively, applies probabilistic clustering to evaluate ward-level vulnerabilities, and utilizes real-time NLP news scraping to validate flood events against ground-truth reports."
    )

    add_heading_1("2. Problem Statement")
    add_body(
        "Current weather warning systems issue general rain alerts but fail to model localized flood risks or severity thresholds. Traditional predictive models evaluate rainfall in isolation, ignoring the cumulative soil moisture saturation caused by antecedent rainfall over the preceding 3-day and 7-day windows. Additionally, disaster management fails to incorporate differences in flooding dynamics between highly urbanized centers (Mumbai City) and regional coastal plains (Konkan Division). There is also a lack of dynamic, probabilistic risk cataloging at the municipal ward level. Finally, unstructured telemetry from social media and digital news feeds remains unintegrated with meteorological systems. This project addresses these limitations by deploying a unified cloud analytics platform with a dual-model predictive pipeline, ward clustering, and automated media scraping."
    )

    add_heading_1("3. Project Objectives")
    add_body("• Develop and implement a dual-model predictive pipeline: (1) An XGBoost Classifier optimized for Mumbai City's localized urban waterlogging, and (2) a Stacking Ensemble (combining Random Forest, Gradient Boosting, and Logistic Regression) for regional forecasting across the Konkan division.")
    add_body("• Incorporate multi-day cumulative antecedent rainfall metrics (3-day and 7-day aggregates) to model soil moisture saturation and absorption limits, reducing target leakage and enhancing predictive accuracy.")
    add_body("• Utilize unsupervised Gaussian Mixture Model (GMM) clustering on municipal wards to calculate probabilistic risk classifications rather than rigid, static boundaries, incorporating features like historical flood hotspots, elevation profiles, and population density.")
    add_body("• Deploy an automated Python scraping and NLP pipeline to monitor digital news, extract flood-related keyword tokens, calculate severity scores, and map historical flood events onto a visual chronological timeline.")
    add_body("• Design and launch a premium glassmorphic dark-theme Single Page Application (SPA) hosted on Vercel with a Supabase PostgreSQL cloud database, rendering real-time predictions and interactive Plotly.js charts.")

    add_heading_1("4. Proposed Methodology & Technical Architecture")
    add_body(
        "The system follows a three-tier architecture ensuring clean separation of concerns between data storage, analytical processing, and client presentation:"
    )
    
    add_heading_2("4.1. Data Storage & Migration (Supabase PostgreSQL)")
    add_body(
        "Meteorological datasets, ward descriptors, and scraped media records are migrated from local formats to a cloud-based Supabase PostgreSQL instance. PostgreSQL provides robust transactional safety, complex queries, and native compatibility with serverless environments. To handle extremely small probabilistic values from GMM clustering computations (e.g., 10^-45), the schema implements DOUBLE PRECISION fields to prevent numerical underflow."
    )

    add_heading_2("4.2. Dual-Model Machine Learning API Layer (FastAPI & Python)")
    add_body(
        "FastAPI serves as the backend engine, running as serverless functions on Vercel. The predictive endpoints load serialized models via pickle to run real-time inference on input parameters (precipitation today, duration, month, and 3-day/7-day cumulative rainfall). The API routes support two distinct model scopes: (1) 'Mumbai Scope' utilizing the localized XGBoost model, and (2) 'Konkan Scope' utilizing the Stacking Ensemble model, returning categorical risk levels (No Flood, Slight, Moderate, Severe) and probability percentages."
    )

    add_heading_2("4.3. Presentation Layer (Tailwind CSS, JS & Plotly.js)")
    add_body(
        "The front-end is designed as a modern, interactive dashboard hosted statically on Vercel. It features a responsive layout styled with Tailwind CSS, utilizing glassmorphic effects and custom animations. Interactive charts are rendered dynamically using Plotly.js, including: (1) A monthly distribution box plot of historical rainfall, (2) a scatter plot correlating daily intensity and 7-day soil saturation, (3) a Pearson correlation matrix heatmap, and (4) an NLP media severity timeline."
    )

    add_heading_1("5. Hardware & Software Requirements")
    add_heading_2("5.1. Software Requirements")
    add_body("Windows 10/11 or Ubuntu Linux 20.04+", bold_prefix="Operating System: ")
    add_body("Python 3.10+, JavaScript (ES6+), HTML5/CSS3", bold_prefix="Programming Languages: ")
    add_body("FastAPI, Uvicorn, Pandas, Scikit-Learn, XGBoost, Psycopg2-Binary", bold_prefix="Backend Framework & ML libraries: ")
    add_body("Tailwind CSS, Lucide Icons, Plotly.js", bold_prefix="Frontend Tech Stack: ")
    add_body("Supabase (PostgreSQL), GitHub, Vercel Serverless Platform", bold_prefix="Database & Deployment: ")

    add_heading_2("5.2. Hardware Requirements")
    add_body("Intel Core i5 / AMD Ryzen 5 processor or higher", bold_prefix="Processor: ")
    add_body("8 GB RAM (16 GB recommended for local ML training)", bold_prefix="Memory: ")
    add_body("256 GB SSD (Solid State Drive) or higher", bold_prefix="Storage: ")
    add_body("Broadband internet connection for cloud database synchronization", bold_prefix="Network: ")

    add_heading_1("6. Expected Outcomes")
    add_body(
        "The project delivers a working, cloud-deployed dashboard powered by a high-accuracy ML model (82.4% test accuracy, 60% flood recall). By providing probabilistic ward-level clustering, a historical correlation exploration tool, and real-time news analytics, the system moves beyond basic weather reports. It offers an actionable tool for municipal officers to simulate scenarios, optimize evacuation pathways, allocate drainage pumps, and manage transit schedules, significantly minimizing monsoon damages across Mumbai."
    )

    # Save documents
    desktop_path = r"C:\Users\User\OneDrive\Desktop\FloodSense_Mumbai_Konkan_Synopsis.docx"
    downloads_path = r"C:\Users\User\Downloads\FloodSense_Mumbai_Konkan_Synopsis.docx"
    
    # Save to desktop
    try:
        doc.save(desktop_path)
        print(f"Successfully saved to Desktop: {desktop_path}")
    except Exception as e:
        print(f"Could not save to Desktop: {e}")

    try:
        doc.save(downloads_path)
        print(f"Successfully saved to Downloads: {downloads_path}")
    except Exception as e:
        print(f"Could not save to Downloads: {e}")

if __name__ == "__main__":
    create_synopsis()
